from docx import Document
from docx.shared import Inches

def add_word(photo_1: str = None, photo_2: str = None, data: dict = None, cheh: str = None, teg: str = None):
    document = Document()

    document.add_heading('Заявка', 0)

    if photo_2 is None and photo_1 is None:
        pass

    else:
        if photo_2 is None:
            document.add_paragraph(
                'Фото проблемы', style='List Bullet'
            )
            document.add_picture(f'img/{photo_1}.jpg', width=Inches(6))
            document.add_page_break()

        elif photo_1 is None:
            document.add_paragraph(
                'Фото решения', style='List Bullet'
            )
            document.add_picture(f'img/{photo_2}.jpg', width=Inches(6))
            document.add_page_break()

        else:
            document.add_paragraph(
                'Фото проблемы', style='List Bullet'
            )
            document.add_picture(f'img/{photo_1}.jpg', width=Inches(6))
            document.add_page_break()
            document.add_paragraph(
                'Фото решения', style='List Bullet'
            )
            document.add_picture(f'img/{photo_2}.jpg', width=Inches(6))
            document.add_page_break()

    # document.add_paragraph(
    #     'Фото проблемы', style='List Bullet'
    # )
    # document.add_picture(f'img/{photo}.jpg', width=Inches(6))
    # document.add_paragraph(
    #     'Фото решения', style='List Bullet'
    # )
    # document.add_picture(f'img/{photo}.jpg', width=Inches(6), height=Inches(3))
    document.add_paragraph(
        data
    )
    document.save(f'word/{cheh + teg}.docx')